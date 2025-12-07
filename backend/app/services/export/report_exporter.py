"""
Report Export Service.

Handles conversion of LLM Markdown outputs to PDF and Word documents.
Uses WeasyPrint for PDF generation and python-docx for Word documents.
Supports optional S3 storage for audit trails and re-downloads.
"""

import os
import io
import datetime
import markdown2
from pathlib import Path
from typing import Optional, Dict, Any, Tuple
import logging

# Template engine
from jinja2 import Environment, FileSystemLoader

# PDF generation
from weasyprint import HTML, CSS

# Word generation
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def _get_storage_service():
    """Get S3 storage service if available."""
    try:
        from app.storage.s3 import get_storage_service
        return get_storage_service()
    except Exception as e:
        logger.debug(f"S3 not available: {e}")
        return None


class ReportExporter:
    """Service for exporting reports to PDF and Word formats."""
    
    def __init__(self):
        self.templates_dir = Path(__file__).parent / "templates"
        self.env = Environment(loader=FileSystemLoader(str(self.templates_dir)))
        self.css_path = self.templates_dir / "styles.css"
        
    def _markdown_to_html(self, markdown_text: str) -> str:
        """Convert Markdown to HTML with extras."""
        return markdown2.markdown(
            markdown_text,
            extras=[
                "fenced-code-blocks",
                "tables",
                "header-ids",
                "break-on-newline",
                "target-blank-links"
            ]
        )
    
    def privacy_report_json_to_markdown(self, report_data: Dict[str, Any], generator_metadata: Dict[str, Any] = None) -> str:
        """
        Convert structured privacy report JSON to comprehensive Markdown.
        
        Args:
            report_data: Privacy report JSON with sections like privacy_budget, risk_assessment, metrics, etc.
            generator_metadata: Generator metadata (name, type, etc.)
            
        Returns:
            Formatted markdown string with all report sections
        """
        lines = []
        
        # Title and metadata
        lines.append("# Privacy Compliance Report")
        lines.append("")
        
        if generator_metadata:
            lines.append("## Generator Information")
            lines.append("")
            if generator_metadata.get("name"):
                lines.append(f"**Generator Name:** {generator_metadata['name']}")
            if generator_metadata.get("type"):
                lines.append(f"**Generator Type:** {generator_metadata['type']}")
            if generator_metadata.get("id"):
                lines.append(f"**Generator ID:** {generator_metadata['id']}")
            lines.append("")
        
        # Privacy Budget Section
        if report_data.get("privacy_budget"):
            budget = report_data["privacy_budget"]
            lines.append("## Privacy Budget")
            lines.append("")
            
            if budget.get("target_epsilon"):
                lines.append(f"**Target Epsilon (ε):** {budget['target_epsilon']}")
            if budget.get("target_delta"):
                lines.append(f"**Target Delta (δ):** {budget['target_delta']}")
            if budget.get("used_epsilon"):
                lines.append(f"**Used Epsilon:** {budget['used_epsilon']}")
            if budget.get("remaining_budget"):
                lines.append(f"**Remaining Budget:** {budget['remaining_budget']}")
            if budget.get("privacy_level"):
                lines.append(f"**Privacy Level:** {budget['privacy_level']}")
            
            lines.append("")
            
            if budget.get("explanation"):
                lines.append("### Explanation")
                lines.append(budget["explanation"])
                lines.append("")
        
        # Privacy Metrics Section
        if report_data.get("privacy_metrics") and isinstance(report_data["privacy_metrics"], list):
            lines.append("## Privacy Metrics")
            lines.append("")
            
            for metric in report_data["privacy_metrics"]:
                if metric.get("name"):
                    lines.append(f"### {metric['name']}")
                if metric.get("value"):
                    lines.append(f"**Value:** {metric['value']}")
                if metric.get("description"):
                    lines.append(f"**Description:** {metric['description']}")
                if metric.get("interpretation"):
                    lines.append(f"**Interpretation:** {metric['interpretation']}")
                lines.append("")
        
        # Risk Assessment Section
        if report_data.get("risk_assessment"):
            risk = report_data["risk_assessment"]
            lines.append("## Risk Assessment")
            lines.append("")
            
            if risk.get("overall_risk"):
                lines.append(f"**Overall Risk Level:** {risk['overall_risk']}")
                lines.append("")
            
            if risk.get("risks") and isinstance(risk["risks"], list):
                lines.append("### Identified Risks")
                lines.append("")
                for r in risk["risks"]:
                    if r.get("risk"):
                        lines.append(f"#### {r['risk']}")
                    if r.get("severity"):
                        lines.append(f"**Severity:** {r['severity']}")
                    if r.get("description"):
                        lines.append(f"**Description:** {r['description']}")
                    if r.get("mitigation"):
                        lines.append(f"**Mitigation:** {r['mitigation']}")
                    lines.append("")
        
        # Recommendations Section
        if report_data.get("recommendations") and isinstance(report_data["recommendations"], list):
            lines.append("## Recommendations")
            lines.append("")
            
            for idx, rec in enumerate(report_data["recommendations"], 1):
                if isinstance(rec, dict):
                    if rec.get("title"):
                        lines.append(f"### {idx}. {rec['title']}")
                    if rec.get("description"):
                        lines.append(rec["description"])
                    if rec.get("priority"):
                        lines.append(f"**Priority:** {rec['priority']}")
                    if rec.get("implementation"):
                        lines.append(f"**Implementation:** {rec['implementation']}")
                else:
                    lines.append(f"{idx}. {rec}")
                lines.append("")
        
        # Compliance Status Section
        if report_data.get("compliance_status"):
            compliance = report_data["compliance_status"]
            lines.append("## Compliance Status")
            lines.append("")
            
            if compliance.get("gdpr"):
                lines.append(f"**GDPR Compliance:** {compliance['gdpr']}")
            if compliance.get("hipaa"):
                lines.append(f"**HIPAA Compliance:** {compliance['hipaa']}")
            if compliance.get("ccpa"):
                lines.append(f"**CCPA Compliance:** {compliance['ccpa']}")
            
            lines.append("")
        
        # Additional Report Content (fallback for generic markdown)
        if report_data.get("report") and isinstance(report_data["report"], str):
            lines.append("## Detailed Analysis")
            lines.append("")
            lines.append(report_data["report"])
            lines.append("")
        
        return "\n".join(lines)

    def export_to_pdf(
        self,
        content_markdown: str,
        title: str,
        metadata: Dict[str, Any] = None,
        output_path: Optional[str] = None
    ) -> bytes:
        """
        Export Markdown content to PDF.
        
        Args:
            content_markdown: The markdown content (e.g., from LLM)
            title: Report title
            metadata: Additional metadata (generator_id, date, etc.)
            output_path: Optional path to save the file
            
        Returns:
            PDF bytes
        """
        try:
            # Convert Markdown to HTML
            html_content = self._markdown_to_html(content_markdown)
            
            # Prepare context for template
            context = {
                "title": title,
                "content": html_content,
                "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "year": datetime.datetime.now().year,
                **(metadata or {})
            }
            
            # Render HTML template
            template = self.env.get_template("base_report.html")
            rendered_html = template.render(**context)
            
            # Generate PDF
            pdf_doc = HTML(string=rendered_html, base_url=str(self.templates_dir))
            pdf_bytes = pdf_doc.write_pdf(stylesheets=[CSS(self.css_path)])
            
            # Save to file if path provided
            if output_path:
                with open(output_path, "wb") as f:
                    f.write(pdf_bytes)
                    
            return pdf_bytes
            
        except Exception as e:
            logger.error(f"Failed to export PDF: {e}")
            raise

    def export_to_docx(
        self,
        content_markdown: str,
        title: str,
        metadata: Dict[str, Any] = None,
        output_path: Optional[str] = None
    ) -> str:
        """
        Export Markdown content to Word document.
        
        Note: This is a simplified conversion. For complex Markdown,
        we parse line by line or use a library like pypandoc (requires system install).
        Here we'll do a basic mapping for headers and paragraphs.
        
        Args:
            content_markdown: The markdown content
            title: Report title
            metadata: Additional metadata
            output_path: Path to save the file
            
        Returns:
            Path to saved file
        """
        try:
            doc = Document()
            
            # Title
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            if metadata:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                run.italic = True
                
                for key, value in metadata.items():
                    if value:
                        run = p.add_run(f"{key.replace('_', ' ').title()}: {value}\n")
                        run.font.size = Pt(9)
            
            doc.add_paragraph()  # Spacer
            
            # Process Markdown content (Basic parser)
            # In a real enterprise app, we might use `pandoc` but that requires system dependencies.
            # We'll stick to a python-only approach for portability.
            
            lines = content_markdown.split('\n')
            code_block = False
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Code blocks
                if line.startswith('```'):
                    code_block = not code_block
                    continue
                    
                if code_block:
                    p = doc.add_paragraph(line)
                    p.style = 'Quote'  # Use Quote style for code
                    p.paragraph_format.left_indent = Inches(0.5)
                    continue
                
                # Headers
                if line.startswith('# '):
                    doc.add_heading(line[2:], 1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], 2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], 3)
                
                # Lists
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line[0].isdigit() and line[1:3] == '. ':
                    doc.add_paragraph(line[3:], style='List Number')
                    
                # Blockquotes
                elif line.startswith('> '):
                    p = doc.add_paragraph(line[2:])
                    p.style = 'Quote'
                    
                # Normal text
                else:
                    doc.add_paragraph(line)
            
            # Footer
            section = doc.sections[0]
            footer = section.footer
            p = footer.paragraphs[0]
            p.text = f"Confidential - Generated by Synthetic Data Studio - {datetime.datetime.now().year}"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save
            if not output_path:
                # Create temp path if none provided
                import tempfile
                fd, output_path = tempfile.mkstemp(suffix=".docx")
                os.close(fd)
                
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to export DOCX: {e}")
            raise

    # ==================== S3 Upload Methods ====================

    def export_pdf_to_s3(
        self,
        content_markdown: str,
        title: str,
        user_id: str,
        export_type: str = "report",
        metadata: Dict[str, Any] = None,
    ) -> Tuple[bytes, Dict[str, Any]]:
        """
        Export Markdown to PDF and upload to S3.
        
        Args:
            content_markdown: Markdown content
            title: Report title
            user_id: User ID for S3 path
            export_type: Type of export (model_card, privacy_report, etc.)
            metadata: Additional metadata
            
        Returns:
            Tuple of (pdf_bytes, s3_info) where s3_info contains:
            - s3_key: The S3 object key
            - s3_bucket: The bucket name
            - file_size_bytes: File size
            - download_url: Presigned download URL
        """
        # Generate PDF
        pdf_bytes = self.export_to_pdf(
            content_markdown=content_markdown,
            title=title,
            metadata=metadata
        )
        
        # Upload to S3
        storage = _get_storage_service()
        if not storage:
            raise RuntimeError("S3 storage not available")
        
        # Create filename
        safe_title = "".join(c if c.isalnum() or c in "-_" else "_" for c in title[:50])
        filename = f"{safe_title}.pdf"
        
        # Upload
        file_obj = io.BytesIO(pdf_bytes)
        upload_result = storage.upload_export(
            file_obj=file_obj,
            user_id=user_id,
            filename=filename,
            export_type=export_type,
            content_type="application/pdf",
            metadata={
                "title": title,
                "export_type": export_type,
                **(metadata or {})
            }
        )
        
        # Generate download URL
        download_url = storage.generate_download_url(
            key=upload_result["key"],
            filename=filename,
            expires_in=3600  # 1 hour
        )
        
        return pdf_bytes, {
            "s3_key": upload_result["key"],
            "s3_bucket": storage.bucket_name,
            "file_size_bytes": len(pdf_bytes),
            "download_url": download_url,
            "filename": filename
        }

    def export_docx_to_s3(
        self,
        content_markdown: str,
        title: str,
        user_id: str,
        export_type: str = "report",
        metadata: Dict[str, Any] = None,
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Export Markdown to DOCX and upload to S3.
        
        Args:
            content_markdown: Markdown content
            title: Report title
            user_id: User ID for S3 path
            export_type: Type of export (model_card, privacy_report, etc.)
            metadata: Additional metadata
            
        Returns:
            Tuple of (local_path, s3_info) where s3_info contains:
            - s3_key: The S3 object key
            - s3_bucket: The bucket name
            - file_size_bytes: File size
            - download_url: Presigned download URL
        """
        # Generate DOCX
        local_path = self.export_to_docx(
            content_markdown=content_markdown,
            title=title,
            metadata=metadata
        )
        
        # Upload to S3
        storage = _get_storage_service()
        if not storage:
            raise RuntimeError("S3 storage not available")
        
        # Create filename
        safe_title = "".join(c if c.isalnum() or c in "-_" else "_" for c in title[:50])
        filename = f"{safe_title}.docx"
        
        # Read file and upload
        with open(local_path, "rb") as f:
            file_size = os.path.getsize(local_path)
            upload_result = storage.upload_export(
                file_obj=f,
                user_id=user_id,
                filename=filename,
                export_type=export_type,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                metadata={
                    "title": title,
                    "export_type": export_type,
                    **(metadata or {})
                }
            )
        
        # Generate download URL
        download_url = storage.generate_download_url(
            key=upload_result["key"],
            filename=filename,
            expires_in=3600  # 1 hour
        )
        
        return local_path, {
            "s3_key": upload_result["key"],
            "s3_bucket": storage.bucket_name,
            "file_size_bytes": file_size,
            "download_url": download_url,
            "filename": filename
        }


# Singleton instance
report_exporter = ReportExporter()
